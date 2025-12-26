
import os
import re
import base64
import requests
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(md_path, docx_path):
    doc = Document()
    
    # helper to add formatted text
    def add_formatted_text(paragraph, text):
        # bold logic: **text**
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    is_mermaid_block = False
    code_lines = []

    for line in lines:
        line = line.strip()
        
        # Code Block Handling
        if line.startswith('```'):
            if in_code_block:
                # End of block
                in_code_block = False
                
                if is_mermaid_block:
                    mermaid_code = '\n'.join(code_lines)
                    try:
                        # Encode for mermaid.ink
                        base64_str = base64.urlsafe_b64encode(mermaid_code.encode('utf-8')).decode('utf-8')
                        url = f"https://mermaid.ink/img/{base64_str}"
                        print(f"Fetching diagram... {url[:50]}...")
                        
                        response = requests.get(url)
                        if response.status_code == 200:
                            image_stream = BytesIO(response.content)
                            p = doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run()
                            run.add_picture(image_stream, width=Inches(6))
                            p = doc.add_paragraph('Figure: System Diagram')
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p.style = 'Caption'
                        else:
                            print(f"Failed to fetch image: {response.status_code}")
                            # Fallback to text
                            p = doc.add_paragraph('[Diagram Fetch Failed - Rendering Code]')
                            p.style = 'No Spacing'
                            run = p.add_run(mermaid_code)
                            run.font.name = 'Courier New'
                            run.font.size = Pt(8)
                    except Exception as e:
                        print(f"Error rendering diagram: {e}")
                        # Fallback
                        p = doc.add_paragraph()
                        p.add_run(mermaid_code)

                    is_mermaid_block = False
                else:
                    # Regular Code Block
                    p = doc.add_paragraph()
                    p.style = 'No Spacing'
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(50, 50, 50)
                
                code_lines = []
            else:
                # Start of block
                in_code_block = True
                if 'mermaid' in line:
                    is_mermaid_block = True
                else:
                    is_mermaid_block = False
            continue
            
        if in_code_block:
            code_lines.append(line)
            continue

        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        
        # Lists
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, line[2:])
        
        # Normal Text
        elif line:
            p = doc.add_paragraph()
            add_formatted_text(p, line)

    try:
        doc.save(docx_path)
        print(f"Successfully saved to {docx_path}")
    except PermissionError:
        print(f"Permission denied for {docx_path}. File might be open. Saving as v2...")
        new_path = docx_path.replace(".docx", "_v2.docx")
        doc.save(new_path)
        print(f"Successfully saved to {new_path}")

if __name__ == "__main__":
    ARTIFACT_DIR = r"C:\Users\USER\.gemini\antigravity\brain\181f9321-c44f-4dd7-843b-6a15814c867a"
    md_file = os.path.join(ARTIFACT_DIR, "system_design.md")
    docx_file = os.path.join(ARTIFACT_DIR, "Unified_MIC_Platform_System_Design.docx")
    
    markdown_to_docx(md_file, docx_file)
